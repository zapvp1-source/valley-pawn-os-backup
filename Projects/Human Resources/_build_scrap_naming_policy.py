#!/usr/bin/env python3
"""Build the Gold Scrap Bucket Naming Standard policy (.docx) — ONE PAGE incl. signature block."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x2D, 0x1A, 0x5E)
BLUE = RGBColor(0x00, 0x99, 0xDD)
GREY = RGBColor(0x59, 0x59, 0x59)

BODY = 7.6
HEAD = 8.8

OUT = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(OUT, "Gold_Scrap_Bucket_Naming_Standard_2026-08.docx")

doc = Document()
s = doc.sections[0]
s.top_margin = Inches(0.34)
s.bottom_margin = Inches(0.24)
s.left_margin = Inches(0.55)
s.right_margin = Inches(0.55)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(BODY)
st.paragraph_format.space_after = Pt(2)
st.paragraph_format.space_before = Pt(0)


def para(text, size=BODY, bold=False, color=None, after=2, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3.5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper())
    r.font.size = Pt(HEAD)
    r.bold = True
    r.font.color.rgb = NAVY
    return p


def bullet(text, indent=0.16):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    r = p.add_run(text)
    r.font.size = Pt(BODY)
    return p


# ---------------- Header ----------------
para("VALLEY PAWN  |  FULL CIRCLE FINANCE INC", size=11, bold=True, color=NAVY, after=0)
para("Policies & Procedures Manual — Operational Procedure", size=BODY, color=BLUE, after=3)
para("GOLD SCRAP BUCKET NAMING STANDARD", size=13, bold=True, color=NAVY, after=1)
para("Effective Date: August 1, 2026     |     Applies To: All Store Locations & Store Personnel     |     "
     "Owner: Market Manager     |     Approved By: Chief Executive Officer",
     size=7.8, color=GREY, after=4)

# ---------------- Purpose ----------------
heading("1. Purpose")
para("Each store has historically named its gold scrap refining buckets differently, with no year identifier, and "
     "identically named buckets recur across calendar years. This makes month-to-month and store-to-store comparison "
     "unreliable and has produced inaccurate reporting. This procedure establishes one required naming format "
     "company-wide so scrap volume can be measured accurately by month, by store, and year over year.")

# ---------------- Policy ----------------
heading("2. Policy")
bullet("2.1  Every gold scrap refining bucket created in the point-of-sale system must be named in this exact format: "
       "YYYY-MM GOLD  (no stones) or  YYYY-MM GOLD WITH STONES  (contains stones).")
bullet("2.2  YYYY-MM is the four-digit year and two-digit month in which the bucket COLLECTS gold — not the month it "
       "is closed. Example: a bucket opened August 1, 2026 is named 2026-08 GOLD.")
bullet("2.3  A separate bucket is required for stone-bearing and non-stone material. The two are never combined.")
bullet("2.4  Buckets are opened on the first day of the collection month and closed and posted the following month, "
       "consistent with existing practice. No change is made to the closing or shipping schedule.")
bullet("2.5  Bucket names may not be abbreviated, reordered, or supplemented with additional wording, initials, "
       "dates, or store identifiers. The format above is the complete name.")
bullet("2.6  Any change to this naming format requires approval by the Chief Executive Officer.")

# ---------------- Procedure ----------------
heading("3. Procedure")
bullet("3.1  On the first business day of each month, the Store Manager opens two scrap refining buckets for that "
       "month using the exact names in Section 2.1.")
bullet("3.2  Store personnel place scrap gold into the correct bucket at the time of intake, separating stone-bearing "
       "material from non-stone material.")
bullet("3.3  In the following month, on the store's normal schedule, the Store Manager verifies the recorded weight is "
       "entered for each bucket, then closes and posts it. A bucket may not be posted with a blank weight.")
bullet("3.4  The Store Manager confirms the bucket name matches the required format before closing. A misnamed bucket "
       "is corrected prior to posting.")

# ---------------- Exceptions ----------------
heading("4. Discrepancy & Exception Handling")
bullet("4.1  If a bucket is discovered misnamed after posting, the Store Manager notifies the Market Manager the same "
       "business day with the store, month, and bucket involved.")
bullet("4.2  If a posted bucket is found to be missing its recorded weight, the Market Manager is notified the same "
       "business day so the figure can be sourced and entered before monthly reporting is finalized.")
bullet("4.3  Items unresolved by the Market Manager are escalated to the Chief Executive Officer the same business day.")
bullet("4.4  Buckets created before the effective date of this procedure are not renamed. This standard applies "
       "going forward only.")

# ---------------- Roles ----------------
heading("5. Roles & Responsibilities")

tbl = doc.add_table(rows=6, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = False
widths = [Inches(1.55), Inches(2.35), Inches(3.35)]
hdr = ["Location", "Store Manager", "Responsibility"]
rows = [
    ("Culpeper", "Sandi Cole", "Opens, names, verifies weight, and closes both monthly buckets."),
    ("Harrisonburg", "Walker Tapley", "Opens, names, verifies weight, and closes both monthly buckets."),
    ("Lexington", "Uriah Tiglao", "Opens, names, verifies weight, and closes both monthly buckets."),
    ("Roanoke", "Benjie Moore", "Opens, names, verifies weight, and closes both monthly buckets."),
    ("Waynesboro", "Chadd McClintic", "Opens, names, verifies weight, and closes both monthly buckets."),
]
for j, h in enumerate(hdr):
    c = tbl.rows[0].cells[j]
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(BODY)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "2D1A5E")
    c._tc.get_or_add_tcPr().append(sh)
for i, row in enumerate(rows, start=1):
    for j, val in enumerate(row):
        c = tbl.rows[i].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val)
        r.font.size = Pt(BODY)
        if j == 0:
            r.bold = True
for row in tbl.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w

para("The Market Manager owns this procedure, monitors compliance across all locations, and reports exceptions to the "
     "Chief Executive Officer.", after=2, before=3)

# ---------------- Relationship ----------------
heading("6. Relationship to Other Company Controls")
para("This procedure supplements, and does not supersede, existing inventory, precious-metals handling, and "
     "loss-prevention controls in the Policies & Procedures Manual and the Employee Handbook. Where an existing "
     "control imposes a stricter requirement, the stricter requirement governs.")

# ---------------- Acknowledgment ----------------
heading("7. Acknowledgment")
para("This procedure is distributed for electronic signature through the Company's payroll and HR system. By signing, "
     "the employee acknowledges receipt and understanding of this procedure and agrees to follow it. The Company may "
     "amend, modify, or discontinue this procedure at any time, with or without notice, at its sole discretion. "
     "Nothing in this procedure creates a contract of employment or alters the at-will employment relationship, under "
     "which either the employee or the Company may end the employment relationship at any time, with or without cause "
     "or notice.", after=3)

# ---------------- Signature block ----------------
sig = doc.add_table(rows=2, cols=4)
sig.alignment = WD_TABLE_ALIGNMENT.LEFT
sig.autofit = False
sw = [Inches(2.6), Inches(1.75), Inches(1.5), Inches(1.4)]
labels = ["Employee Signature (via Gusto)", "Printed Name", "Store / Location", "Date"]
for j, lab in enumerate(labels):
    c = sig.rows[0].cells[j]
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(7)
    r = p.add_run("_" * (30 if j == 0 else 20))
    r.font.size = Pt(BODY)
    c2 = sig.rows[1].cells[j]
    c2.text = ""
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(lab)
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = GREY
for row in sig.rows:
    for j, w in enumerate(sw):
        row.cells[j].width = w

doc.save(DOCX)
print("wrote", DOCX)
